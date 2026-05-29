"""Claude API summarization + .docx export."""

from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

CLAUDE_MODEL = "claude-sonnet-4-20250514"


def _client(api_key: str):
    from app.core.claude_api import make_client
    return make_client(api_key)


def summarize_part(
    transcript_text: str,
    speakers_reference: Dict[str, Any],
    summary_prompt: str,
    api_key: str,
    part_number: int = 1,
) -> str:
    """Run a single transcript part through Claude with the user's chosen prompt."""
    client = _client(api_key)
    campaign_context_block = json.dumps(
        {
            "campaign": speakers_reference.get("campaign", ""),
            "context": speakers_reference.get("context", ""),
            "players": speakers_reference.get("players", []),
        },
        indent=2,
        ensure_ascii=False,
    )

    full_prompt = (
        f"{summary_prompt}\n\n"
        f"========================================================\n"
        f"CAMPAIGN CONTEXT (from speakers.json):\n"
        f"========================================================\n"
        f"{campaign_context_block}\n\n"
        f"========================================================\n"
        f"TRANSCRIPT — PART {part_number}\n"
        f"========================================================\n"
        f"{transcript_text}\n"
    )

    resp = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=4000,
        messages=[{"role": "user", "content": full_prompt}],
    )
    return resp.content[0].text


def consolidate_summaries(
    part_summaries: List[str],
    speakers_reference: Dict[str, Any],
    api_key: str,
) -> Dict[str, str]:
    """Consolidate per-part summaries into a unified session summary."""
    client = _client(api_key)
    parts_block = "\n\n".join(
        f"--- PART {i + 1} ---\n{summary}" for i, summary in enumerate(part_summaries)
    )
    context_block = json.dumps(
        {
            "campaign": speakers_reference.get("campaign", ""),
            "context": speakers_reference.get("context", ""),
        },
        indent=2,
        ensure_ascii=False,
    )

    prompt = (
        "You are consolidating individual part summaries from a D&D session into one "
        "unified session summary.\n\n"
        f"CAMPAIGN CONTEXT:\n{context_block}\n\n"
        f"INDIVIDUAL PART SUMMARIES:\n{parts_block}\n\n"
        "Produce a single unified session summary following the same structure as the "
        "individual summaries but synthesized across all parts. De-duplicate events, "
        "consolidate loot, and produce a clean 'OPEN THREADS GOING INTO NEXT SESSION' "
        "section.\n\n"
        "Begin your response with a single line in the form:\n"
        "SESSION NAME: <thematic session name here>\n\n"
        "Then proceed with the summary sections."
    )
    resp = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=4000,
        messages=[{"role": "user", "content": prompt}],
    )
    text = resp.content[0].text
    name_match = re.search(r"^\s*SESSION NAME:\s*(.+)$", text, flags=re.MULTILINE)
    session_name = name_match.group(1).strip() if name_match else "Session Summary"
    body = text[name_match.end():].lstrip() if name_match else text
    return {"session_name": session_name, "body": body, "raw": text}


def safe_filename(name: str) -> str:
    name = name.strip()
    name = re.sub(r"[^A-Za-z0-9 _\-]", "", name)
    name = re.sub(r"\s+", "_", name)
    return name or "Session_Summary"


def write_docx(
    output_path: str,
    session_name: str,
    consolidated_body: str,
    part_summaries: List[str],
    campaign_name: str = "",
    model_used: str = CLAUDE_MODEL,
) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(session_name, level=1)

    meta = doc.add_paragraph()
    run = meta.add_run(
        f"Campaign: {campaign_name or 'N/A'} | Generated: "
        f"{datetime.now().strftime('%Y-%m-%d %H:%M')} | Model: {model_used}"
    )
    run.italic = True
    run.font.size = Pt(9)

    _render_summary_body(doc, consolidated_body)

    if part_summaries:
        doc.add_heading("Individual Part Summaries", level=2)
        for i, summary in enumerate(part_summaries, start=1):
            doc.add_heading(f"Part {i}", level=3)
            for line in summary.splitlines():
                doc.add_paragraph(line)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))


def _render_summary_body(doc, body: str) -> None:
    """Convert the Claude response body into Word headings, paragraphs and bullets."""
    lines = body.splitlines()
    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        # Markdown-style headings
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        # ALL CAPS section headers (e.g. KEY STORY DEVELOPMENTS)
        elif (re.match(r"^[A-Z][A-Z &/'\-,()0-9]{4,}$", stripped)
              and not stripped.startswith(("- ", "* ", "•"))):
            doc.add_heading(stripped, level=2)
        # Bullet point
        elif stripped.startswith(("- ", "* ", "• ")):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif stripped.startswith("•"):
            doc.add_paragraph(stripped[1:].strip(), style="List Bullet")
        # Table rows -> just plain paragraphs (good enough for this app)
        else:
            doc.add_paragraph(stripped)


def parse_session_name_from_text(text: str) -> Optional[str]:
    m = re.search(r"^\s*SESSION NAME:\s*(.+)$", text, flags=re.MULTILINE)
    return m.group(1).strip() if m else None
